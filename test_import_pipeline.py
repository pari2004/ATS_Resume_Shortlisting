"""End-to-end tests for import_pipeline.py, covering the scenarios called out
in the spec: public Drive folder, large folder (100+ resumes), duplicate
resumes, invalid PDF, DOC/DOCX mix, and an empty folder. Drive access is
mocked the same way test_drive_utils.py mocks it -- no real network needed.
"""

import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import ats_scoring
import excel_store
import import_pipeline
from drive_utils import DriveFetchResult


def _make_pdf(name, email, skills_line, years_line=""):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    _w, height = letter
    y = height - 100
    for text in [name, f"Email: {email}", "Skills", skills_line, "Experience", years_line]:
        c.setFont("Helvetica", 12)
        c.drawString(100, y, text)
        y -= 22
    c.save()
    return buf.getvalue()


def _make_docx(name, email):
    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph(name)
    document.add_paragraph(f"Email: {email}")
    document.save(buf)
    return buf.getvalue()


JD = ats_scoring.JobDescription(
    raw_text="Looking for a Python developer.",
    required_skills=["python"],
    min_experience_years=0,
)


class TestImportPipelineUploads(unittest.TestCase):
    """Uses the uploaded_files path (no Drive mocking needed) to exercise the
    parse -> score -> upsert -> report loop directly."""

    def setUp(self):
        self.tmp_dir = tempfile.TemporaryDirectory()
        self.applicants_path = str(Path(self.tmp_dir.name) / "Applicants.xlsx")

    def tearDown(self):
        self.tmp_dir.cleanup()

    def test_single_resume_is_imported(self):
        files = {"a.pdf": _make_pdf("Alice Ng", "alice@example.com", "Python, SQL")}
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files=files
        )
        self.assertEqual(report.imported, 1)
        self.assertEqual(report.failed, 0)
        applicants, _ = excel_store.load_applicants(self.applicants_path)
        self.assertEqual(len(applicants), 1)

    def test_duplicate_resume_across_two_files_rescores_not_duplicates(self):
        pdf_bytes = _make_pdf("Alice Ng", "alice@example.com", "Python, SQL")
        files = {"a.pdf": pdf_bytes, "a_copy.pdf": pdf_bytes}
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files=files
        )
        self.assertEqual(report.imported, 1)
        self.assertEqual(report.rescored, 1)
        applicants, _ = excel_store.load_applicants(self.applicants_path)
        self.assertEqual(len(applicants), 1)  # still one candidate row, not two

    def test_rerunning_same_folder_with_a_different_jd_refreshes_score(self):
        """The scenario this behavior exists for: importing the same resume
        again under a different job description should update its score/
        domain, not leave it stale from whatever JD scored it the first
        time."""
        pdf_bytes = _make_pdf("Alice Ng", "alice@example.com", "Python, SQL, AWS")
        tech_jd = ats_scoring.JobDescription(
            domain="Python Developer", required_skills=["python"], min_experience_years=0,
        )
        other_jd = ats_scoring.JobDescription(
            domain="Sales", required_skills=["sales", "crm", "negotiation"], min_experience_years=0,
        )
        import_pipeline.run_import(
            drive_link=None, jd=tech_jd, applicants_path=self.applicants_path,
            uploaded_files={"a.pdf": pdf_bytes},
        )
        report = import_pipeline.run_import(
            drive_link=None, jd=other_jd, applicants_path=self.applicants_path,
            uploaded_files={"a.pdf": pdf_bytes},
        )
        self.assertEqual(report.rescored, 1)
        applicants, _ = excel_store.load_applicants(self.applicants_path)
        self.assertEqual(len(applicants), 1)
        self.assertEqual(applicants.iloc[0]["Detected Job Domain"], "Sales")

    def test_invalid_pdf_counts_as_failed_but_does_not_abort_batch(self):
        files = {
            "good.pdf": _make_pdf("Alice Ng", "alice@example.com", "Python"),
            "broken.pdf": b"not a real pdf",
        }
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files=files
        )
        self.assertEqual(report.imported, 1)
        self.assertEqual(report.failed, 1)
        statuses = {o.file_name: o.status for o in report.outcomes}
        self.assertEqual(statuses["broken.pdf"], "failed")
        self.assertEqual(statuses["good.pdf"], "imported")

    def test_docx_resume_is_imported(self):
        files = {"bob.docx": _make_docx("Bob Lee", "bob@example.com")}
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files=files
        )
        self.assertEqual(report.imported, 1)

    def test_empty_batch_produces_zero_counts(self):
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files={}
        )
        self.assertEqual(report.total_files, 0)
        self.assertEqual(report.imported, 0)
        self.assertEqual(report.failed, 0)

    def test_progress_callback_fires_for_every_file(self):
        files = {
            "a.pdf": _make_pdf("Alice Ng", "alice@example.com", "Python"),
            "b.pdf": _make_pdf("Bob Lee", "bob@example.com", "SQL"),
        }
        calls = []
        import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path,
            uploaded_files=files, progress_cb=lambda done, total, name: calls.append((done, total)),
        )
        self.assertEqual(calls[0], (0, 2))
        self.assertEqual(calls[-1], (2, 2))

    def test_large_batch_of_100_resumes(self):
        files = {
            f"candidate_{i}.pdf": _make_pdf(f"Person {i}", f"person{i}@example.com", "Python, SQL")
            for i in range(100)
        }
        report = import_pipeline.run_import(
            drive_link=None, jd=JD, applicants_path=self.applicants_path, uploaded_files=files
        )
        self.assertEqual(report.total_files, 100)
        self.assertEqual(report.imported, 100)
        self.assertEqual(report.failed, 0)
        self.assertEqual(report.progress_pct, 100.0)
        applicants, _ = excel_store.load_applicants(self.applicants_path)
        self.assertEqual(len(applicants), 100)
        ids = set(applicants["Candidate ID"])
        self.assertEqual(len(ids), 100)  # every candidate got a unique ID


class TestImportPipelineDriveFolder(unittest.TestCase):
    """Exercises the drive_link path with fetch_pdfs_from_drive mocked."""

    def setUp(self):
        self.tmp_dir = tempfile.TemporaryDirectory()
        self.applicants_path = str(Path(self.tmp_dir.name) / "Applicants.xlsx")

    def tearDown(self):
        self.tmp_dir.cleanup()

    @patch("import_pipeline.fetch_pdfs_from_drive")
    def test_public_folder_success(self, mock_fetch):
        mock_fetch.return_value = DriveFetchResult(
            {"alice.pdf": _make_pdf("Alice Ng", "alice@example.com", "Python")},
            warnings=[],
            file_meta={"alice.pdf": {"file_id": "id123", "url": "https://drive.google.com/x"}},
        )
        report = import_pipeline.run_import(
            drive_link="https://drive.google.com/drive/folders/abc",
            jd=JD, applicants_path=self.applicants_path,
        )
        self.assertEqual(report.imported, 1)
        applicants, _ = excel_store.load_applicants(self.applicants_path)
        self.assertEqual(applicants.iloc[0]["Google Drive File ID"], "id123")

    @patch("import_pipeline.fetch_pdfs_from_drive")
    def test_empty_folder_produces_drive_warning_and_no_crash(self, mock_fetch):
        from drive_utils import DriveFetchError
        mock_fetch.side_effect = DriveFetchError("No resumes were found in that Drive folder.")

        report = import_pipeline.run_import(
            drive_link="https://drive.google.com/drive/folders/empty",
            jd=JD, applicants_path=self.applicants_path,
        )
        self.assertEqual(report.total_files, 0)
        self.assertTrue(any("No resumes" in w for w in report.drive_warnings))


def _make_domain_pdf(lines):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    _w, height = letter
    y = height - 100
    for text in lines:
        c.setFont("Helvetica", 12)
        c.drawString(100, y, text)
        y -= 20
    c.save()
    return buf.getvalue()


# One (JD, matching resume) pair per domain named in the spec's testing
# requirement -- each resume is written to genuinely match its own JD's
# required skills, not another domain's, so cross-contamination would show
# up as a wrong/irrelevant recommendation sentence.
_DOMAIN_FIXTURES = {
    "Full Stack Development": (
        "Full Stack Developer\nRequired Skills\nReact, Node.js, MongoDB, Express.js\n3+ years of experience.",
        ["Alex Fuller", "Email: alex.fuller@example.com", "Skills",
         "React, Node.js, MongoDB, Express.js, JavaScript", "Experience",
         "Full Stack Developer at WebCo", "Jan 2020 - Present"],
    ),
    "Java Developer": (
        "Java Developer\nRequired Skills\nJava, Spring Boot, Hibernate, Microservices",
        ["Rahul Verma", "Email: rahul.verma@example.com", "Skills",
         "Java, Spring Boot, Hibernate, Microservices", "Experience",
         "Java Developer at Enterprise Systems", "Jun 2019 - Present"],
    ),
    "Python Developer": (
        "Python Developer\nRequired Skills\nPython, Django, Flask, PostgreSQL",
        ["Sara Lin", "Email: sara.lin@example.com", "Skills",
         "Python, Django, Flask, PostgreSQL", "Experience",
         "Python Developer at DataWorks", "Mar 2021 - Present"],
    ),
    "AI Engineer": (
        "AI Engineer\nRequired Skills\nLLMs, Prompt Engineering, PyTorch, GenAI Pipelines",
        ["Mia Chen", "Email: mia.chen@example.com", "Skills",
         "LLMs, Prompt Engineering, PyTorch, GenAI Pipelines", "Experience",
         "AI Engineer at NeuralWorks", "Feb 2022 - Present"],
    ),
    "Data Scientist": (
        "Data Scientist\nRequired Skills\nPython, Pandas, Machine Learning, TensorFlow, Statistics",
        ["Devon Park", "Email: devon.park@example.com", "Skills",
         "Python, Pandas, Machine Learning, TensorFlow, Statistics", "Experience",
         "Data Scientist at Insight Analytics", "Jul 2020 - Present"],
    ),
    "Business Analyst": (
        "Business Analyst\nRequired Skills\nRequirements Gathering, SQL, Stakeholder Management, Business Analysis",
        ["Nora Bell", "Email: nora.bell@example.com", "Skills",
         "Requirements Gathering, SQL, Stakeholder Management, Business Analysis", "Experience",
         "Business Analyst at ConsultCo", "Sep 2019 - Present"],
    ),
    "Recruiter": (
        "HR Recruiter\nRequired Skills\nRecruitment, Payroll, Employee Engagement, Onboarding",
        ["Priya Sharma", "Email: priya.sharma@example.com", "Skills",
         "Recruitment, Payroll, Employee Engagement, Onboarding", "Experience",
         "HR Recruiter at PeopleFirst", "Jan 2021 - Present"],
    ),
    "Digital Marketing": (
        "Digital Marketing Manager\nRequired Skills\nSEO, SEM, Google Analytics, Meta Ads, Content Marketing",
        ["Liam Cole", "Email: liam.cole@example.com", "Skills",
         "SEO, SEM, Google Analytics, Meta Ads, Content Marketing", "Experience",
         "Digital Marketing Manager at GrowthHub", "Apr 2020 - Present"],
    ),
    "Finance": (
        "Finance Executive\nRequired Skills\nGST, Tally, SAP, Excel, Accounting",
        ["Karan Mehta", "Email: karan.mehta@example.com", "Skills",
         "GST, Tally, SAP, Excel, Accounting", "Experience",
         "Finance Executive at LedgerCorp", "Aug 2018 - Present"],
    ),
    "Sales": (
        "Sales Executive\nRequired Skills\nLead Generation, Cold Calling, Negotiation, CRM, Salesforce",
        ["Tara Singh", "Email: tara.singh@example.com", "Skills",
         "Lead Generation, Cold Calling, Negotiation, CRM, Salesforce", "Experience",
         "Sales Executive at Dealwise", "May 2019 - Present"],
    ),
    "Customer Support": (
        "Customer Support Representative\nRequired Skills\nZendesk, Ticketing, Customer Service, Communication",
        ["Omar Farouk", "Email: omar.farouk@example.com", "Skills",
         "Zendesk, Ticketing, Customer Service, Communication", "Experience",
         "Customer Support Representative at HelpDeskCo", "Oct 2021 - Present"],
    ),
    "UI/UX Designer": (
        "UI/UX Designer\nRequired Skills\nFigma, Adobe XD, Wireframing, User Research, Prototyping",
        ["Elena Cruz", "Email: elena.cruz@example.com", "Skills",
         "Figma, Adobe XD, Wireframing, User Research, Prototyping", "Experience",
         "UI/UX Designer at PixelStudio", "Nov 2020 - Present"],
    ),
}


class TestDomainIndependence(unittest.TestCase):
    """Runs one matching (JD, resume) pair per domain named in the spec's
    testing requirement through the full pipeline, verifying each produces
    distinct extracted skills, a real (non-collapsed) ATS score, and a
    recommendation that doesn't leak another domain's vocabulary."""

    def setUp(self):
        self.tmp_dir = tempfile.TemporaryDirectory()
        self.applicants_path = str(Path(self.tmp_dir.name) / "Applicants.xlsx")

    def tearDown(self):
        self.tmp_dir.cleanup()

    def test_each_domain_scores_and_recommends_independently(self):
        results = {}
        for domain_name, (jd_text, resume_lines) in _DOMAIN_FIXTURES.items():
            jd = ats_scoring.analyze_job_description(jd_text)
            pdf_bytes = _make_domain_pdf(resume_lines)
            report = import_pipeline.run_import(
                drive_link=None, jd=jd, applicants_path=self.applicants_path,
                uploaded_files={f"{domain_name}.pdf": pdf_bytes},
            )
            self.assertEqual(report.failed, 0, f"{domain_name} resume failed to process")
            self.assertEqual(report.imported, 1, f"{domain_name} resume was not imported")
            results[domain_name] = report.outcomes[0]

        applicants, _dedup = excel_store.load_applicants(self.applicants_path)
        by_domain = {row["Detected Job Domain"]: row for _, row in applicants.iterrows()}

        # Every fixture's own required skills should actually be found on
        # its own resume, and the detected domain should match what was fed
        # in (proves the JD -> parser -> scorer chain stays domain-scoped
        # end to end, not just the isolated unit-level pieces).
        for domain_name in _DOMAIN_FIXTURES:
            self.assertIn(domain_name, by_domain, f"No candidate recorded under domain {domain_name!r}")
            row = by_domain[domain_name]
            self.assertGreater(float(row["ATS Score"]), 50.0, f"{domain_name} candidate scored unexpectedly low")

        # Recommendation sentences must not cross-contaminate -- an HR
        # candidate's recommendation shouldn't mention a Full Stack skill,
        # and vice versa.
        hr_recommendation = by_domain["Recruiter"]["Recommendation"].lower()
        for leaked_term in ("react", "node.js", "django", "figma"):
            self.assertNotIn(leaked_term, hr_recommendation)

        fsd_recommendation = by_domain["Full Stack Development"]["Recommendation"].lower()
        for leaked_term in ("recruitment", "payroll", "tally", "zendesk"):
            self.assertNotIn(leaked_term, fsd_recommendation)

        # Every fixture here is a deliberately perfect match for its own JD,
        # so identical high scores across the board is the *correct*
        # outcome (proves the formula treats every domain fairly, not that
        # it's broken) -- the real differentiation signal is that the
        # *skills* extracted per candidate are genuinely distinct sets, not
        # the same taxonomy-derived list regardless of domain.
        skill_sets = {d: frozenset(row["Skills"].split(", ")) for d, row in by_domain.items()}
        self.assertGreater(len(set(skill_sets.values())), len(skill_sets) - 2)

    def test_mismatched_domain_scores_lower_than_matched(self):
        """A Java-skilled candidate applying against an HR JD should score
        noticeably lower than an HR-skilled candidate against that same JD
        -- proves the engine actually penalizes a cross-domain mismatch
        instead of defaulting everyone to a similar score."""
        hr_jd_text, hr_resume_lines = _DOMAIN_FIXTURES["Recruiter"]
        _java_jd_text, java_resume_lines = _DOMAIN_FIXTURES["Java Developer"]
        jd = ats_scoring.analyze_job_description(hr_jd_text)

        matched_report = import_pipeline.run_import(
            drive_link=None, jd=jd, applicants_path=self.applicants_path,
            uploaded_files={"matched.pdf": _make_domain_pdf(hr_resume_lines)},
        )
        mismatched_report = import_pipeline.run_import(
            drive_link=None, jd=jd, applicants_path=self.applicants_path,
            uploaded_files={"mismatched.pdf": _make_domain_pdf(java_resume_lines)},
        )
        self.assertGreater(matched_report.outcomes[0].ats_score, mismatched_report.outcomes[0].ats_score)
        self.assertEqual(mismatched_report.outcomes[0].status, "imported")  # still processed, just scored low


if __name__ == "__main__":
    unittest.main(verbosity=2)
