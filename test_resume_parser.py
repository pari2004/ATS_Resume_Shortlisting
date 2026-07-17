"""Tests for resume_parser.py: text extraction (PDF/DOCX/DOC, including
corrupt/invalid files) and structured field parsing."""

import io
import tempfile
import unittest
from pathlib import Path

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import resume_parser


def _make_pdf_bytes(lines):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    _width, height = letter
    y = height - 100
    for font, size, text in lines:
        c.setFont(font, size)
        c.drawString(100, y, text)
        y -= 22
    c.save()
    return buf.getvalue()


FULL_RESUME_LINES = [
    ("Helvetica-Bold", 18, "Jane A Smith"),
    ("Helvetica", 12, "Email: jane.smith@example.com | Phone: +1 415-555-2671"),
    ("Helvetica", 12, "LinkedIn: linkedin.com/in/janesmith  GitHub: github.com/janesmith"),
    ("Helvetica", 12, "Portfolio: janesmith.dev"),
    ("Helvetica-Bold", 14, "Skills"),
    ("Helvetica", 12, "Python, SQL, AWS, Docker, React"),
    ("Helvetica-Bold", 14, "Experience"),
    ("Helvetica", 12, "Lead Engineer at Initech"),
    ("Helvetica", 12, "Mar 2021 - Present"),
    ("Helvetica", 12, "Engineer at Globex"),
    ("Helvetica", 12, "Jan 2018 - Feb 2021"),
    ("Helvetica-Bold", 14, "Education"),
    ("Helvetica", 12, "M.S. Computer Science, Tech University, 2018"),
    ("Helvetica-Bold", 14, "Certifications"),
    ("Helvetica", 12, "AWS Certified Developer"),
    ("Helvetica-Bold", 14, "Projects"),
    ("Helvetica", 12, "Built an internal analytics dashboard"),
]


class TestPdfTextExtraction(unittest.TestCase):
    def test_extract_text_from_valid_pdf(self):
        pdf_bytes = _make_pdf_bytes(FULL_RESUME_LINES)
        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(pdf_bytes, "jane.pdf", Path(tmp))
        self.assertTrue(ok)
        self.assertFalse(is_scanned)
        self.assertIn("Jane A Smith", text)

    def test_invalid_pdf_bytes_fails_gracefully(self):
        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(
                b"this is not a real pdf file", "broken.pdf", Path(tmp)
            )
        self.assertFalse(ok)
        self.assertEqual(text, "")

    def test_unsupported_extension_returns_not_ok(self):
        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(b"hello", "notes.txt", Path(tmp))
        self.assertFalse(ok)


class TestDocxTextExtraction(unittest.TestCase):
    def test_extract_text_from_docx(self):
        import docx

        buf = io.BytesIO()
        document = docx.Document()
        document.add_paragraph("Bob Builder")
        document.add_paragraph("Email: bob@example.com")
        document.save(buf)

        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(buf.getvalue(), "bob.docx", Path(tmp))
        self.assertTrue(ok)
        self.assertIn("Bob Builder", text)
        self.assertIn("bob@example.com", text)

    def test_legacy_doc_without_usable_text_fails_gracefully(self):
        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(
                b"\x00\x01\x02\x03binary-garbage\x04\x05", "old.doc", Path(tmp)
            )
        self.assertFalse(ok)


class TestParseResume(unittest.TestCase):
    def setUp(self):
        pdf_bytes = _make_pdf_bytes(FULL_RESUME_LINES)
        with tempfile.TemporaryDirectory() as tmp:
            text, is_scanned, ok = resume_parser.extract_text(pdf_bytes, "jane.pdf", Path(tmp))
        self.assertTrue(ok)
        self.parsed = resume_parser.parse_resume(text, "jane.pdf", is_scanned=is_scanned)

    def test_name(self):
        self.assertEqual(self.parsed.name, "Jane A Smith")

    def test_email(self):
        self.assertEqual(self.parsed.email, "jane.smith@example.com")

    def test_phone_extracted(self):
        self.assertTrue(self.parsed.phone)

    def test_linkedin_and_github(self):
        self.assertIn("linkedin.com/in/janesmith", self.parsed.linkedin)
        self.assertIn("github.com/janesmith", self.parsed.github)

    def test_skills_found(self):
        for skill in ["python", "sql", "aws", "docker", "react"]:
            self.assertIn(skill, self.parsed.skills)

    def test_non_taxonomy_skill_found_verbatim_in_skills_section(self):
        text = "Someone\nsomeone@example.com\nSkills\nGST, Tally, Bank Reconciliation\n"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertIn("gst", parsed.skills)
        self.assertIn("tally", parsed.skills)
        self.assertIn("bank reconciliation", parsed.skills)

    def test_languages_section_parsed(self):
        text = "Someone\nsomeone@example.com\nLanguages\nEnglish\nHindi\nSpanish\n"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.languages, ["English", "Hindi", "Spanish"])

    def test_achievements_section_parsed(self):
        text = "Someone\nsomeone@example.com\nAchievements\nEmployee of the Month, 2024\nTop Sales Performer\n"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertTrue(any("Employee of the Month" in a for a in parsed.achievements))

    def test_education_found(self):
        self.assertTrue(any("Tech University" in e for e in self.parsed.education))

    def test_certifications_found(self):
        self.assertTrue(any("AWS Certified" in c for c in self.parsed.certifications))

    def test_total_experience_computed(self):
        self.assertGreater(self.parsed.total_experience_years, 0)

    def test_current_role(self):
        self.assertIn("Initech", self.parsed.current_company)
        self.assertIn("Lead Engineer", self.parsed.current_designation)

    def test_resume_hash_deterministic(self):
        h1 = resume_parser.compute_resume_hash("Some Resume Text")
        h2 = resume_parser.compute_resume_hash("some   resume text")
        self.assertEqual(h1, h2)  # whitespace/case-insensitive normalization

    def test_missing_fields_default_gracefully(self):
        parsed = resume_parser.parse_resume("Just some unrelated text with no structure.", "x.pdf")
        self.assertEqual(parsed.email, "")
        self.assertEqual(parsed.skills, [])
        self.assertEqual(parsed.total_experience_years, 0.0)


class TestNameExtractionEdgeCases(unittest.TestCase):
    """Regression tests for real garbled-name cases seen on production data:
    job-title fragments folded into filenames, and section headers like
    'PROFESSIONAL EXPERIENCE' being mistaken for an all-caps name."""

    def test_filename_with_job_title_fragments_does_not_pollute_name(self):
        text = "someone@example.com\n\nSkills\nReact, Node.js"
        parsed = resume_parser.parse_resume(
            text, "Aarif_FullStackDeveloper_ReactDeveloper.pdf"
        )
        self.assertNotIn("FullStackDeveloper", parsed.name)
        self.assertNotIn("Developer", parsed.name)

    def test_filename_job_title_abbreviation_is_dropped(self):
        text = "someone@example.com"
        parsed = resume_parser.parse_resume(text, "Ansh_Gupta_SE.pdf")
        self.assertEqual(parsed.name, "Ansh Gupta")

    def test_section_header_not_mistaken_for_allcaps_name(self):
        text = "\n".join([
            "someone@example.com",
            "PROFESSIONAL EXPERIENCE",
            "Software Engineer at Acme",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertNotEqual(parsed.name, "PROFESSIONAL EXPERIENCE")

    def test_work_experience_header_not_mistaken_for_name(self):
        text = "\n".join([
            "someone@example.com",
            "WORK EXPERIENCE",
            "Senior Developer at Beta Corp",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertNotEqual(parsed.name, "WORK EXPERIENCE")

    def test_genuine_allcaps_name_is_still_recognized(self):
        text = "\n".join([
            "JOHN MICHAEL DOE",
            "john.doe@example.com",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "JOHN MICHAEL DOE")

    def test_titlecase_name_after_section_header_is_found(self):
        text = "\n".join([
            "CURRICULUM VITAE",
            "Rohan Kushwah",
            "rohan@example.com",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "Rohan Kushwah")

    def test_name_sharing_a_line_with_email_is_still_found(self):
        text = "+91 7000000000\nASHISH KUMAR someone@example.com\nSUMMARY:"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "ASHISH KUMAR")

    def test_cid_icon_font_artifacts_are_stripped_before_name_detection(self):
        text = "RAVI SAHAY (cid:131) +91-7000000000\nEmail: ravi@example.com"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "RAVI SAHAY")

    def test_mixed_case_initials_name_is_recognized(self):
        text = "MD Kaif Manzar\nSoftware Development Engineer\nkaif@example.com"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "MD Kaif Manzar")

    def test_dotted_initials_name_is_recognized(self):
        text = "M.V. KUNDAPPAN\nFull Stack Developer\nmv@example.com"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "M.V. KUNDAPPAN")

    def test_name_split_across_two_header_lines_is_joined(self):
        text = "AFTAB\nKHAN\nWeb developer\n+91 6200000000 someone@example.com"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "AFTAB KHAN")

    def test_name_with_job_title_glued_on_same_line_is_trimmed(self):
        text = (
            "Pankaj Yadav Full Stack Developer | MERN Stack | React.js\n"
            "py@example.com 8000000000 Noida"
        )
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "Pankaj Yadav")

    def test_first_name_only_does_not_absorb_an_address_line(self):
        text = "PRIYANSHU\nBhiwadi , Rajasthan - 301019\n+91-8000000000 someone@example.com"
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertEqual(parsed.name, "PRIYANSHU")

    def test_filename_single_letter_initial_is_accepted(self):
        parsed = resume_parser.parse_resume("someone@example.com", "Aditya_S_resume.pdf")
        self.assertEqual(parsed.name, "Aditya S")

    def test_glued_camelcase_name_stays_unknown_rather_than_guessing(self):
        # Indistinguishable by shape from a glued job title like
        # "FullStackDeveloper" -- correctly refuses to guess.
        text = "RohitBande\nFrontend Developer (ReactJS & NextJS)\nsomeone@example.com"
        parsed = resume_parser.parse_resume(text, "Resume_2026_july.pdf")
        self.assertEqual(parsed.name, "Unknown")


class TestExperienceExtractionEdgeCases(unittest.TestCase):
    """Regression tests for real gaps found on non-tech (BDA/sales) resumes:
    numeric MM/YYYY date ranges and bare duration phrases with no calendar
    dates at all, both of which previously computed to 0.0 years."""

    def test_numeric_month_year_date_range_is_parsed(self):
        text = "\n".join([
            "Someone Person",
            "Internship",
            "ABC Company - Accountant",
            "Patna - 06/2025 - 07/2025",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertGreater(parsed.total_experience_years, 0)

    def test_numeric_date_range_to_present_is_parsed(self):
        text = "\n".join([
            "Someone Person",
            "Experience",
            "Company X - Analyst",
            "03/2023 - Present",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertGreater(parsed.total_experience_years, 0)

    def test_bare_months_duration_with_no_dates_is_parsed(self):
        text = "\n".join([
            "Someone Person",
            "Experience",
            "Accounts Intern, Some Stadium (5 months)",
            "- Assisted in filing and documentation.",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertAlmostEqual(parsed.total_experience_years, 0.4, delta=0.05)

    def test_bare_duration_without_parens_is_parsed(self):
        text = "\n".join([
            "Someone Person",
            "Internship Experience",
            "Company Name: Project Title: Duration:",
            "Some Company Trainee 2 months",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertAlmostEqual(parsed.total_experience_years, 0.2, delta=0.05)

    def test_explicit_years_and_months_duration_is_parsed(self):
        text = "\n".join([
            "Someone Person",
            "Experience",
            "Worked as an analyst for 1 year 3 months at Some Firm.",
        ])
        parsed = resume_parser.parse_resume(text, "resume.pdf")
        self.assertAlmostEqual(parsed.total_experience_years, 1.25, delta=0.06)


if __name__ == "__main__":
    unittest.main(verbosity=2)
